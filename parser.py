import re
import os
import spacy
import pdfplumber
from docx import Document

# Lazy loading of spaCy model with automatic download
_nlp = None

def get_nlp():
    global _nlp
    if _nlp is None:
        try:
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Download if not present
            print("spaCy model 'en_core_web_sm' not found. Downloading...")
            os.system("python -m spacy download en_core_web_sm")
            _nlp = spacy.load("en_core_web_sm")
    return _nlp

# Common skill keywords
SKILL_DB = [
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "php", "go", "rust", "swift", "kotlin", "scala", "html", "css", "sql", "bash", "shell",
    # Frameworks / Libraries
    "react", "angular", "vue", "node.js", "express", "flask", "django", "spring boot", "laravel", "asp.net", "next.js", "bootstrap", "tailwind", "jquery",
    "spacy", "nltk", "scikit-learn", "tensorflow", "pytorch", "keras", "pandas", "numpy", "opencv", "flask-sqlalchemy", "fastapi",
    # Databases
    "postgresql", "mysql", "sqlite", "mongodb", "redis", "cassandra", "mariadb", "oracle", "sql server", "dynamodb",
    # DevOps / Cloud
    "docker", "kubernetes", "aws", "azure", "gcp", "google cloud", "git", "github", "gitlab", "jenkins", "terraform", "ansible", "ci/cd", "linux", "nginx",
    # Concepts / Other
    "machine learning", "deep learning", "nlp", "natural language processing", "computer vision", "data science", "data engineering", "agile", "scrum", "project management",
    "rest api", "graphql", "microservices", "system design", "oop", "algorithms", "data structures", "cloud computing"
]

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text]
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"Error reading DOCX {docx_path}: {e}")
    return text

def extract_contact_info(text):
    # Email regex
    email_regex = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    emails = re.findall(email_regex, text)
    email = emails[0] if emails else None

    # Phone regex (matches various formats like +1 123 456 7890, (123) 456-7890, 123-456-7890)
    phone_regex = r'(?:(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})'
    phones = re.findall(phone_regex, text)
    phone = phones[0] if phones else None

    return email, phone

def extract_name(text):
    nlp = get_nlp()
    # Name is typically at the top. Let's look at the first few lines of text.
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    first_few_lines = lines[:4]
    
    # Try parsing the first few lines with spaCy NER to check for PERSON
    for line in first_few_lines:
        # Ignore contact info or email lines
        if '@' in line or any(char.isdigit() for char in line if char in '0123456789') and len(line) < 20:
            continue
        doc = nlp(line)
        for ent in doc.ents:
            if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                # Basic sanity check: name shouldn't contain common noise
                name = ent.text.strip()
                # Clean up punctuation/trailing noise
                name = re.sub(r'[^\w\s-]', '', name)
                return name
                
    # Fallback to the first non-empty line that has 2-3 words and looks like a name
    for line in first_few_lines:
        words = line.split()
        if 2 <= len(words) <= 3 and all(w[0].isupper() for w in words if w.isalpha()):
            # Check for email/phone markers
            if '@' not in line and not any(k in line.lower() for k in ['resume', 'curriculum', 'cv', 'email', 'phone']):
                return line
                
    return "Unknown Candidate"

def extract_skills(text):
    text_lower = text.lower()
    skills_found = []
    for skill in SKILL_DB:
        # Escape skill for regex and search with word boundaries
        # For skills with special characters like c++, c# or node.js, we adjust word boundaries
        escaped_skill = re.escape(skill)
        if skill in ['c++', 'c#', '.net']:
            # Custom boundary for special char skills
            pattern = r'(?:^|[\s,;])' + escaped_skill + r'(?:$|[\s,;])'
        else:
            pattern = r'\b' + escaped_skill + r'\b'
            
        if re.search(pattern, text_lower):
            # Normalize display name
            display_name = skill
            if skill == 'python': display_name = 'Python'
            elif skill == 'javascript': display_name = 'JavaScript'
            elif skill == 'typescript': display_name = 'TypeScript'
            elif skill == 'java': display_name = 'Java'
            elif skill == 'c++': display_name = 'C++'
            elif skill == 'c#': display_name = 'C#'
            elif skill == 'html': display_name = 'HTML'
            elif skill == 'css': display_name = 'CSS'
            elif skill == 'sql': display_name = 'SQL'
            elif skill == 'react': display_name = 'React'
            elif skill == 'angular': display_name = 'Angular'
            elif skill == 'vue': display_name = 'Vue'
            elif skill == 'node.js': display_name = 'Node.js'
            elif skill == 'flask': display_name = 'Flask'
            elif skill == 'django': display_name = 'Django'
            elif skill == 'postgresql': display_name = 'PostgreSQL'
            elif skill == 'mysql': display_name = 'MySQL'
            elif skill == 'sqlite': display_name = 'SQLite'
            elif skill == 'mongodb': display_name = 'MongoDB'
            elif skill == 'redis': display_name = 'Redis'
            elif skill == 'docker': display_name = 'Docker'
            elif skill == 'kubernetes': display_name = 'Kubernetes'
            elif skill == 'aws': display_name = 'AWS'
            elif skill == 'azure': display_name = 'Azure'
            elif skill == 'git': display_name = 'Git'
            elif skill == 'github': display_name = 'GitHub'
            elif skill == 'spacy': display_name = 'spaCy'
            elif skill == 'nltk': display_name = 'NLTK'
            elif skill == 'pytorch': display_name = 'PyTorch'
            elif skill == 'tensorflow': display_name = 'TensorFlow'
            elif skill == 'scikit-learn': display_name = 'Scikit-learn'
            elif skill == 'pandas': display_name = 'Pandas'
            elif skill == 'numpy': display_name = 'NumPy'
            elif skill == 'fastapi': display_name = 'FastAPI'
            else:
                display_name = display_name.title()
            skills_found.append(display_name)
            
    return sorted(list(set(skills_found)))

def parse_sections(text):
    """
    Split text into sections based on headers like Education, Experience, Summary.
    """
    sections = {
        'summary': [],
        'education': [],
        'experience': [],
        'other': []
    }
    
    lines = text.split('\n')
    current_section = 'other'
    
    # Headers indicators
    summary_headers = ['summary', 'profile', 'objective', 'about me', 'professional summary', 'career objective']
    edu_headers = ['education', 'academic', 'qualifications', 'education background', 'degrees']
    exp_headers = ['experience', 'work history', 'employment', 'professional experience', 'work experience', 'history']
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Check if line looks like a header (short, and matches key terms)
        line_lower = line_clean.lower().replace(':', '')
        is_header = False
        
        if len(line_clean) < 40:
            if any(h == line_lower or line_lower.startswith(h + ' ') or line_lower.endswith(' ' + h) for h in summary_headers):
                current_section = 'summary'
                is_header = True
            elif any(h == line_lower or line_lower.startswith(h + ' ') or line_lower.endswith(' ' + h) for h in edu_headers):
                current_section = 'education'
                is_header = True
            elif any(h == line_lower or line_lower.startswith(h + ' ') or line_lower.endswith(' ' + h) for h in exp_headers):
                current_section = 'experience'
                is_header = True
                
        if not is_header:
            sections[current_section].append(line_clean)
            
    return sections

def extract_education(edu_lines):
    education_entries = []
    
    # Degree patterns
    degree_patterns = [
        (r'\bB\.?S\.?\b|\bBachelor\b|\bB\.?Tech\b|\bB\.?E\.?\b|\bB\.?A\.?\b|\bB\.?C\.?A\.?\b', 'Bachelor\'s'),
        (r'\bM\.?S\.?\b|\bMaster\b|\bM\.?Tech\b|\bM\.?B\.?A\.?\b|\bM\.?C\.?A\.?\b', 'Master\'s'),
        (r'\bPh\.?D\.?\b|\bDoctor\b', 'Ph.D.'),
        (r'\bAssociate\b|\bA\.?S\.?\b', 'Associate\'s')
    ]
    
    # Find school names
    school_keywords = [
        'university', 'college', 'school', 'institute', 'academy', 'polytechnic'
    ]
    
    # Group lines that are likely part of the same education entry
    # (usually education entries are separated by a line or close to each other)
    current_entry = {}
    
    for line in edu_lines:
        # Check if line contains a degree
        found_degree = None
        for pattern, deg_name in degree_patterns:
            if re.search(pattern, line, re.IGNORECASE):
                found_degree = deg_name
                # Also try to extract field of study if in the same line
                # e.g., "Bachelor of Science in Computer Science"
                field_match = re.search(r'(?:in|of)\s+([A-Za-z\s]{3,40})(?:\s+from|\s+at|\s*$,;]|\s*\()', line, re.IGNORECASE)
                if field_match:
                    current_entry['field_of_study'] = field_match.group(1).strip()
                break
                
        # Check if line contains a school name
        found_school = None
        if any(keyword in line.lower() for keyword in school_keywords):
            # Extract the school name (usually from start of line or around keywords)
            # A simple regex to capture the name containing the keyword
            school_match = re.search(r'([A-Za-z\s,\.-]+(?:' + '|'.join(school_keywords) + r')[A-Za-z\s,\.-]*)', line, re.IGNORECASE)
            if school_match:
                found_school = school_match.group(1).strip()
                
        # Extract years (start and end)
        years = re.findall(r'\b(?:19|20)\d{2}\b', line)
        
        # If we have a degree or school, start a new entry or update the current one
        if found_degree or found_school:
            if current_entry and (found_degree and 'degree' in current_entry or found_school and 'institution' in current_entry):
                # Flush previous entry if it looks like we're starting a new one
                education_entries.append(current_entry)
                current_entry = {}
                
            if found_degree:
                current_entry['degree'] = found_degree
            if found_school:
                current_entry['institution'] = found_school
                
        if years and current_entry:
            if len(years) >= 2:
                current_entry['start_year'] = years[0]
                current_entry['end_year'] = years[1]
            else:
                if re.search(r'\bpresent\b', line, re.IGNORECASE):
                    current_entry['start_year'] = years[0]
                    current_entry['end_year'] = 'Present'
                else:
                    current_entry['end_year'] = years[0]
            
    if current_entry:
        education_entries.append(current_entry)
        
    # Ensure minimum fields
    parsed_edu = []
    for entry in education_entries:
        if 'degree' in entry or 'institution' in entry:
            parsed_edu.append({
                'degree': entry.get('degree', 'Degree'),
                'institution': entry.get('institution', 'University/College'),
                'field_of_study': entry.get('field_of_study', 'General Studies'),
                'start_year': entry.get('start_year', ''),
                'end_year': entry.get('end_year', '')
            })
            
    return parsed_edu

def extract_experience(exp_lines):
    experience_entries = []
    current_entry = {}
    desc_lines = []
    
    # Common job titles
    job_titles = [
        'developer', 'engineer', 'architect', 'analyst', 'manager', 'consultant', 
        'lead', 'director', 'intern', 'specialist', 'designer', 'administrator',
        'programmer', 'scientist', 'technician', 'officer', 'representative'
    ]
    
    for line in exp_lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        
        # Check if line looks like a job title + company line
        is_job_line = False
        found_title = None
        
        if len(line_clean) < 100:
            # Check for job title
            for title in job_titles:
                if re.search(r'\b' + title + r'\b', line_clean, re.IGNORECASE):
                    found_title = line_clean
                    is_job_line = True
                    break
                    
        # Extract dates (e.g. "Jun 2018 - Present" or "2019 - 2021")
        date_pattern = r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?[a-z]*\.?\s*(?:19|20)\d{2}\s*(?:[-–]|to)\s*(?:Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?[a-z]*\.?\s*(?:19|20)\d{2})\b|\b(?:19|20)\d{2}\s*[-–]\s*(?:Present|(?:19|20)\d{2})\b'
        date_match = re.search(date_pattern, line_clean, re.IGNORECASE)
        
        if is_job_line:
            # If we already have a job title in current_entry, flush it
            if current_entry and 'job_title' in current_entry:
                current_entry['description'] = "\n".join(desc_lines).strip()
                experience_entries.append(current_entry)
                current_entry = {}
                desc_lines = []
                
            # Parse title and company out of the line
            parts = re.split(r'\s+at\s+|\s*-\s*|\s*\|\s*', found_title, flags=re.IGNORECASE)
            if len(parts) >= 2:
                current_entry['job_title'] = parts[0].strip()
                current_entry['company'] = parts[1].strip()
            else:
                current_entry['job_title'] = found_title
                current_entry['company'] = 'Company'
                
            if date_match:
                date_str = date_match.group(0)
                date_parts = re.split(r'[-–]|to', date_str, flags=re.IGNORECASE)
                if len(date_parts) == 2:
                    current_entry['start_date'] = date_parts[0].strip()
                    current_entry['end_date'] = date_parts[1].strip()
                else:
                    current_entry['start_date'] = date_str.strip()
                    current_entry['end_date'] = ''
        elif date_match:
            # Add dates to current entry if not already set
            if current_entry and 'start_date' not in current_entry:
                date_str = date_match.group(0)
                date_parts = re.split(r'[-–]|to', date_str, flags=re.IGNORECASE)
                if len(date_parts) == 2:
                    current_entry['start_date'] = date_parts[0].strip()
                    current_entry['end_date'] = date_parts[1].strip()
                else:
                    current_entry['start_date'] = date_str.strip()
                    current_entry['end_date'] = ''
            else:
                desc_lines.append(line_clean)
        else:
            # Add to description lines
            if current_entry:
                desc_lines.append(line_clean)
                    
    if current_entry:
        current_entry['description'] = "\n".join(desc_lines).strip()
        experience_entries.append(current_entry)
        
    # Standardize
    parsed_exp = []
    for entry in experience_entries:
        parsed_exp.append({
            'job_title': entry.get('job_title', 'Role'),
            'company': entry.get('company', 'Company'),
            'start_date': entry.get('start_date', ''),
            'end_date': entry.get('end_date', ''),
            'description': entry.get('description', '')
        })
        
    return parsed_exp

def parse_resume(file_path):
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        text = ""
        
    if not text.strip():
        return {
            'name': 'Unknown Candidate',
            'email': 'N/A',
            'phone': 'N/A',
            'skills': [],
            'education': [],
            'experience': [],
            'summary': 'Could not extract text from file.',
            'raw_text': ''
        }
        
    # Extract details
    email, phone = extract_contact_info(text)
    name = extract_name(text)
    skills = extract_skills(text)
    
    sections = parse_sections(text)
    
    # Process summary
    summary = "\n".join(sections['summary']).strip()
    if not summary and sections['other']:
        # Fallback to the first few lines of "other" section if no summary section
        summary = "\n".join(sections['other'][:3]).strip()
        
    # Trim summary if too long
    if len(summary) > 500:
        summary = summary[:500] + "..."
        
    education = extract_education(sections['education'])
    # If no education found in education section, try to search the whole text
    if not education:
        education = extract_education(text.split('\n'))
        
    experience = extract_experience(sections['experience'])
    # If no experience found, try looking for titles in the whole text
    if not experience:
        experience = extract_experience(text.split('\n'))
        
    return {
        'name': name,
        'email': email or 'N/A',
        'phone': phone or 'N/A',
        'skills': skills,
        'education': education,
        'experience': experience,
        'summary': summary or 'No professional summary found.',
        'raw_text': text
    }
